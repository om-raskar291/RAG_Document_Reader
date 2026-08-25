import io
import os
import hashlib

import streamlit as st

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_google_genai import (
    GoogleGenerativeAIEmbeddings,
    ChatGoogleGenerativeAI
)

from langchain_chroma import Chroma

from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough


# ============================================================
# PAGE CONFIGURATION
# ============================================================

st.set_page_config(
    page_title="RAG Document Reader",
    page_icon="📚",
    layout="wide"
)


# ============================================================
# API KEY
# ============================================================

GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]

os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY


# ============================================================
# APPLICATION TITLE
# ============================================================

st.title("📚 RAG Document Reader")

st.write(
    "Upload a TXT, PDF, or DOCX document and ask questions "
    "based on its contents."
)


# ============================================================
# FILE UPLOAD
# ============================================================

st.subheader("📂 Upload Your Document")

uploaded_file = st.file_uploader(
    "Drop your file here",
    type=["txt", "pdf", "docx"],
    help="Supported file types: TXT, PDF and DOCX"
)


# ============================================================
# DOCUMENT EXTRACTION
# ============================================================

def extract_text_from_file(uploaded_file):

    file_name = uploaded_file.name

    file_extension = file_name.lower().split(".")[-1]

    file_bytes = uploaded_file.getvalue()


    # --------------------------------------------------------
    # TXT
    # --------------------------------------------------------

    if file_extension == "txt":

        text = file_bytes.decode(
            "utf-8",
            errors="ignore"
        )

        documents = [
            Document(
                page_content=text,
                metadata={
                    "source": file_name
                }
            )
        ]

        return documents


    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    elif file_extension == "pdf":

        pdf_reader = PdfReader(
            io.BytesIO(file_bytes)
        )

        documents = []

        for page_number, page in enumerate(
            pdf_reader.pages
        ):

            text = page.extract_text() or ""

            if text.strip():

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file_name,
                            "page": page_number + 1
                        }
                    )
                )

        return documents


    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    elif file_extension == "docx":

        docx_file = DocxDocument(
            io.BytesIO(file_bytes)
        )

        paragraphs = []

        for paragraph in docx_file.paragraphs:

            if paragraph.text.strip():

                paragraphs.append(
                    paragraph.text
                )

        text = "\n\n".join(paragraphs)

        documents = [
            Document(
                page_content=text,
                metadata={
                    "source": file_name
                }
            )
        ]

        return documents


    else:

        raise ValueError(
            "Unsupported file format."
        )


# ============================================================
# CREATE RAG CHAIN
# ============================================================

@st.cache_resource(show_spinner=False)
def create_rag_chain(
    file_bytes,
    file_name
):

    # --------------------------------------------------------
    # Extract document text
    # --------------------------------------------------------

    class UploadedFileWrapper:

        def __init__(
            self,
            name,
            data
        ):

            self.name = name
            self._data = data

        def getvalue(self):

            return self._data


    uploaded_file = UploadedFileWrapper(
        file_name,
        file_bytes
    )

    documents = extract_text_from_file(
        uploaded_file
    )


    # --------------------------------------------------------
    # Validate extracted text
    # --------------------------------------------------------

    if not documents:

        raise ValueError(
            "No readable text was found in the uploaded file."
        )


    # --------------------------------------------------------
    # Text Splitting
    # --------------------------------------------------------

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )

    splits = splitter.split_documents(
        documents
    )


    # --------------------------------------------------------
    # Gemini Embeddings
    # --------------------------------------------------------

    embeddings = GoogleGenerativeAIEmbeddings(
        model="gemini-embedding-2"
    )


    # --------------------------------------------------------
    # Create unique collection name
    # --------------------------------------------------------

    file_hash = hashlib.md5(
        file_bytes
    ).hexdigest()[:12]

    collection_name = (
        f"rag_document_{file_hash}"
    )


    # --------------------------------------------------------
    # ChromaDB
    # --------------------------------------------------------

    vectorstore = Chroma.from_documents(
        documents=splits,
        embedding=embeddings,
        collection_name=collection_name
    )


    # --------------------------------------------------------
    # Retriever
    # --------------------------------------------------------

    retriever = vectorstore.as_retriever(
        search_kwargs={
            "k": 4
        }
    )


    # --------------------------------------------------------
    # Gemini LLM
    # --------------------------------------------------------

    llm = ChatGoogleGenerativeAI(
        model="gemini-3.5-flash",
        temperature=0.5
    )


    # --------------------------------------------------------
    # Prompt Template
    # --------------------------------------------------------

    template = """
Answer the question based only on the following context:

{context}

Question:
{question}

Instructions:

1. Answer only using the provided context.
2. Do not use outside knowledge.
3. If the answer is not present in the context,
   clearly state that the information is not available
   in the uploaded document.
"""

    prompt = PromptTemplate.from_template(
        template
    )


    # --------------------------------------------------------
    # Format Retrieved Documents
    # --------------------------------------------------------

    def format_docs(docs):

        return "\n\n".join(
            doc.page_content
            for doc in docs
        )


    # --------------------------------------------------------
    # LCEL RAG Chain
    # --------------------------------------------------------

    rag_chain = (
        {
            "context": retriever | format_docs,
            "question": RunnablePassthrough()
        }
        | prompt
        | llm
        | StrOutputParser()
    )


    return (
        rag_chain,
        retriever,
        documents,
        splits
    )


# ============================================================
# PROCESS UPLOADED FILE
# ============================================================

if uploaded_file is not None:

    file_bytes = uploaded_file.getvalue()

    file_name = uploaded_file.name


    # --------------------------------------------------------
    # File information
    # --------------------------------------------------------

    file_size_kb = len(file_bytes) / 1024

    file_extension = (
        file_name.split(".")[-1]
        .upper()
    )


    st.success(
        f"✅ Document uploaded successfully: {file_name}"
    )


    col1, col2, col3 = st.columns(3)


    with col1:

        st.metric(
            "File Type",
            file_extension
        )


    with col2:

        st.metric(
            "File Size",
            f"{file_size_kb:.2f} KB"
        )


    with col3:

        st.metric(
            "Status",
            "Ready"
        )


    # --------------------------------------------------------
    # Create RAG
    # --------------------------------------------------------

    with st.spinner(
        "Processing document and creating vector embeddings..."
    ):

        try:

            (
                rag_chain,
                retriever,
                documents,
                splits
            ) = create_rag_chain(
                file_bytes,
                file_name
            )

        except Exception as e:

            st.error(
                "Unable to process the uploaded document."
            )

            st.exception(e)

            st.stop()


    st.success(
        f"✅ Document processed successfully! "
        f"Created {len(splits)} text chunks."
    )


    # ========================================================
    # DOCUMENT PREVIEW
    # ========================================================

    st.subheader("📖 Document Preview")

    full_text = "\n\n".join(
        document.page_content
        for document in documents
    )


    with st.expander(
        "View extracted document text"
    ):

        st.text_area(
            "Extracted Text",
            full_text,
            height=300
        )


    st.divider()


    # ========================================================
    # QUESTION ANSWERING
    # ========================================================

    st.subheader("💬 Ask Questions About Your Document")


    question = st.text_input(
        "Enter your question:",
        placeholder="Example: Who killed Harry's parents?"
    )


    if st.button(
        "🔍 Get Answer",
        type="primary"
    ):

        if not question.strip():

            st.warning(
                "Please enter a question."
            )

        else:

            with st.spinner(
                "Searching the document..."
            ):

                try:

                    # ----------------------------------------
                    # Retrieve relevant documents
                    # ----------------------------------------

                    retrieved_docs = retriever.invoke(
                        question
                    )


                    # ----------------------------------------
                    # Generate answer
                    # ----------------------------------------

                    answer = rag_chain.invoke(
                        question
                    )


                    # ----------------------------------------
                    # Display answer
                    # ----------------------------------------

                    st.subheader(
                        "🤖 Answer"
                    )

                    st.success(
                        answer
                    )


                    # ----------------------------------------
                    # Display retrieved context
                    # ----------------------------------------

                    st.subheader(
                        "📚 Retrieved Context"
                    )

                    for index, doc in enumerate(
                        retrieved_docs,
                        start=1
                    ):

                        with st.expander(
                            f"Retrieved Document {index}"
                        ):

                            st.write(
                                doc.page_content
                            )

                            if doc.metadata:

                                st.caption(
                                    f"Metadata: {doc.metadata}"
                                )


                except Exception as e:

                    st.error(
                        "An error occurred while "
                        "generating the answer."
                    )

                    st.exception(e)


else:

    # ========================================================
    # NO FILE UPLOADED
    # ========================================================

    st.info(
        "👆 Upload a TXT, PDF, or DOCX document "
        "to start asking questions."
    )


# ============================================================
# FOOTER
# ============================================================

st.divider()

st.caption(
    "Built using Python, LangChain, Google Gemini, "
    "ChromaDB and Streamlit."
)
